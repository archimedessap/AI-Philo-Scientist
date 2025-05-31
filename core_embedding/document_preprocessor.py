#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文档预处理器

预处理学术文献，识别结构，分割文档，过滤无关内容。
支持PDF和Word文档格式。
"""

import os
import re
import json
import math
from typing import List, Dict, Any, Tuple

class DocumentPreprocessor:
    def __init__(self, output_dir="data/preprocessed", max_tokens=50000):
        """
        初始化文档预处理器
        
        Args:
            output_dir: 预处理后文档保存的目录
            max_tokens: 每个文档片段的最大token数量
        """
        self.output_dir = output_dir
        self.max_tokens = max_tokens
        os.makedirs(output_dir, exist_ok=True)
    
    def preprocess_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        预处理文件，根据文件类型调用对应的处理方法
        
        Args:
            file_path: 文件路径
            
        Returns:
            List[Dict[str, Any]]: 分割后的文档段落
        """
        if not os.path.exists(file_path):
            print(f"[ERROR] 文件不存在: {file_path}")
            return []
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.preprocess_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.preprocess_word(file_path)
        else:
            print(f"[ERROR] 不支持的文件类型: {file_ext}")
            return []
    
    def preprocess_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """预处理PDF文件，返回分割后的文档段落"""
        try:
            # 尝试PyPDF2
            from PyPDF2 import PdfReader
            
            reader = PdfReader(pdf_path)
            filename = os.path.basename(pdf_path)
            
            print(f"[INFO] 开始处理PDF: {filename}，共{len(reader.pages)}页")
            
            # 提取全文
            full_text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:  # 避免添加空页面
                    full_text += page_text + "\n\n"
                
            # 记录原始文本长度
            original_length = len(full_text)
            print(f"[INFO] 提取的原始文本长度: {original_length}字符")
            
            # 处理提取的文本
            sections = self._process_extracted_text(full_text, filename)
            
            # 检查处理后的总内容长度
            processed_length = sum(len(section["content"]) for section in sections)
            print(f"[INFO] 处理后的文本总长度: {processed_length}字符 ({processed_length/original_length:.1%})")
            
            return sections
            
        except ImportError:
            try:
                # 尝试pypdf
                from pypdf import PdfReader
                
                reader = PdfReader(pdf_path)
                filename = os.path.basename(pdf_path)
                
                print(f"[INFO] 开始处理PDF: {filename}，共{len(reader.pages)}页")
                
                # 提取全文
                full_text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # 避免添加空页面
                        full_text += page_text + "\n\n"
                
                # 记录原始文本长度
                original_length = len(full_text)
                print(f"[INFO] 提取的原始文本长度: {original_length}字符")
                
                # 处理提取的文本
                sections = self._process_extracted_text(full_text, filename)
                
                # 检查处理后的总内容长度
                processed_length = sum(len(section["content"]) for section in sections)
                print(f"[INFO] 处理后的文本总长度: {processed_length}字符 ({processed_length/original_length:.1%})")
                
                return sections
                
            except ImportError:
                # 创建一个空的PdfReader类，允许代码继续运行
                class PdfReader:
                    def __init__(self, pdf_file):
                        self.pages = []
                        print(f"[MOCK] 创建了模拟的PdfReader对象: {pdf_file}")
        except Exception as e:
            print(f"[ERROR] 处理PDF {pdf_path} 时出错: {str(e)}")
            return []
    
    def preprocess_word(self, word_path: str) -> List[Dict[str, Any]]:
        """预处理Word文件，返回分割后的文档段落"""
        try:
            # 尝试导入docx
            import docx
            
            filename = os.path.basename(word_path)
            print(f"[INFO] 开始处理Word文档: {filename}")
            
            # 读取Word文档
            doc = docx.Document(word_path)
            
            # 提取全文
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            full_text = "\n\n".join(paragraphs)
            
            # 记录原始文本长度
            original_length = len(full_text)
            print(f"[INFO] 提取的原始文本长度: {original_length}字符")
            
            # 处理提取的文本
            sections = self._process_extracted_text(full_text, filename)
            
            # 检查处理后的总内容长度
            processed_length = sum(len(section["content"]) for section in sections)
            print(f"[INFO] 处理后的文本总长度: {processed_length}字符 ({processed_length/original_length:.1%})")
            
            return sections
            
        except ImportError:
            try:
                # 尝试从python-docx导入
                from docx import Document
                # 创建一个模拟的docx模块
                class docx:
                    @staticmethod
                    def Document(file_path):
                        return Document(file_path)
            except ImportError:
                # 创建一个模拟的docx模块
                class docx:
                    @staticmethod
                    def Document(file_path):
                        print(f"[MOCK] 创建了模拟的Document对象: {file_path}")
                        return MockDocument()
                
                class MockDocument:
                    def __init__(self):
                        self.paragraphs = []
                        print("[MOCK] 创建了模拟的Document实例")
        except Exception as e:
            print(f"[ERROR] 处理Word文档 {word_path} 时出错: {str(e)}")
            return []
    
    def _process_extracted_text(self, text: str, filename: str) -> List[Dict[str, Any]]:
        """处理提取的文本，进行分割和保存"""
        # 移除摘要（如果存在且能明确识别）
        abstract_end = self._find_abstract_end(text)
        if abstract_end > 0:
            abstract = text[:abstract_end].strip()
            main_text = text[abstract_end:].strip()
            print(f"[INFO] 已识别并分离摘要部分，摘要长度: {len(abstract)}字符")
        else:
            abstract = ""
            main_text = text
        
        # 过滤参考文献
        refs_start = self._find_references_start(main_text)
        if refs_start > 0:
            refs_text = main_text[refs_start:].strip()
            main_text = main_text[:refs_start].strip()
            print(f"[INFO] 已过滤参考文献部分，参考文献长度: {len(refs_text)}字符")
        
        # 分离主要章节
        sections = self._identify_chapters(main_text)
        
        # 如果找不到明确的章节，将整个正文作为单个部分
        if not sections:
            print("[INFO] 未识别出明确的章节结构，将使用整篇文章作为单个部分并按需分割")
            sections = [{
                "title": "Main Content",
                "content": main_text,
                "type": "main_content",
                "estimated_tokens": self._estimate_tokens(main_text)
            }]
        
        # 如果提取了摘要，将其作为独立部分
        if abstract:
            sections.insert(0, {
                "title": "Abstract",
                "content": abstract,
                "type": "abstract",
                "estimated_tokens": self._estimate_tokens(abstract)
            })
        
        # 优化分割策略，最大化利用token上限
        optimized_sections = self._optimize_sections(sections)
        
        # 保存预处理结果
        output_path = os.path.join(self.output_dir, f"{os.path.splitext(filename)[0]}.json")
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(optimized_sections, f, ensure_ascii=False, indent=2)
            
        print(f"[INFO] 文档 {filename} 已预处理并分为{len(optimized_sections)}个部分，保存到 {output_path}")
        return optimized_sections
    
    def _find_abstract_end(self, text: str) -> int:
        """找到摘要部分的结束位置"""
        abstract_patterns = [
            r"(?i)abstract\s*\n",
            r"(?i)摘要\s*\n"
        ]
        
        intro_patterns = [
            r"(?i)introduction\s*\n", 
            r"(?i)1[\.\s]+introduction",
            r"(?i)引言\s*\n",
            r"(?i)导言\s*\n",
            r"(?i)1[\.\s]+引言"
        ]
        
        # 首先检查是否有摘要标记
        abstract_match = None
        for pattern in abstract_patterns:
            match = re.search(pattern, text[:500])  # 只在文档开头寻找
            if match:
                abstract_match = match
                break
                
        if not abstract_match:
            return 0  # 没有明确的摘要部分
            
        # 然后寻找引言开始的位置
        for pattern in intro_patterns:
            match = re.search(pattern, text[abstract_match.end():2000])
            if match:
                return abstract_match.end() + match.start()
                
        # 如果找不到引言，尝试找第一个明显的段落分隔
        paragraphs = text[abstract_match.end():].split("\n\n")
        if len(paragraphs) > 1:
            return abstract_match.end() + len(paragraphs[0]) + 2
            
        return 0  # 无法确定摘要结束位置
    
    def _find_references_start(self, text: str) -> int:
        """找到参考文献部分的开始位置"""
        ref_patterns = [
            r"(?i)references\s*\n",
            r"(?i)bibliography\s*\n",
            r"(?i)引用文献\s*\n",
            r"(?i)参考文献\s*\n"
        ]
        
        # 在文本后半部分寻找参考文献标记
        half_point = len(text) // 2
        for pattern in ref_patterns:
            match = re.search(pattern, text[half_point:])
            if match:
                return half_point + match.start()
                
        # 特殊情况：查找可能的参考文献格式（如[1] Author...)
        ref_format = r"\[\d+\]\s+[A-Z][a-z]+,*\s+[A-Z]\."
        matches = list(re.finditer(ref_format, text))
        if matches and len(matches) > 3:  # 确保有多个连续的参考文献
            # 检查最后几个匹配是否接近文章末尾且彼此接近
            last_matches = matches[-5:] if len(matches) >= 5 else matches
            pos = last_matches[0].start()
            
            # 确认这确实是参考文献部分（检查前面的文本是否有类似"References"的内容）
            preceding_text = text[max(0, pos-100):pos]
            if re.search(r"(?i)(reference|bibliography|cited|文献)", preceding_text):
                return pos
        
        return 0  # 找不到参考文献部分
    
    def _identify_chapters(self, text: str) -> List[Dict[str, Any]]:
        """识别文档的章节结构，尽量识别所有可能的章节"""
        
        # 识别多种可能的章节标记模式
        chapter_patterns = [
            # 编号章节模式（如"1. Introduction", "第一章 引言"）
            r"(?:第?(?:[0-9一二三四五六七八九十]+)[章部节\.\s]+)([^\n]+)",
            r"(?:[0-9]+\.?[0-9]*\s+)([A-Z][^\n]+)",  # 1. Title 或 1.1 Title
            r"(?:[IVX]+\.?\s+)([A-Z][^\n]+)",        # 罗马数字编号
            
            # 特定章节名称
            r"(?i)(introduction|background|methodology|methods|results|discussion|conclusion|framework|theory)",
            r"(引言|背景|方法|结果|讨论|结论|框架|理论)"
        ]
        
        # 找出所有可能的章节起始位置
        chapter_matches = []
        for pattern in chapter_patterns:
            for match in re.finditer(pattern, text):
                title = match.group(1) if match.lastindex else match.group(0)
                chapter_matches.append((match.start(), title.strip()))
        
        # 移除重复的章节（可能被多个模式匹配）
        filtered_matches = []
        last_pos = -1
        for pos, title in sorted(chapter_matches):
            if pos > last_pos + 20:  # 避免非常接近的匹配
                filtered_matches.append((pos, title))
                last_pos = pos
        
        # 分割文档为章节
        sections = []
        for i, (start, title) in enumerate(filtered_matches):
            end = filtered_matches[i+1][0] if i < len(filtered_matches)-1 else len(text)
            section_text = text[start:end].strip()
            
            # 确保章节有足够内容
            if len(section_text) > 200:  # 太短的可能是误识别
                sections.append({
                    "title": title,
                    "content": section_text,
                    "type": "chapter",
                    "estimated_tokens": self._estimate_tokens(section_text)
                })
        
        print(f"[INFO] 识别出{len(sections)}个章节")
        return sections
    
    def _optimize_sections(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        优化章节分割，尽量保持章节完整性，并最大化利用token上限
        """
        optimized_sections = []
        current_batch = []
        current_tokens = 0
        
        # 对于token数已经超过限制的大章节，需要单独分割
        oversized_sections = []
        normal_sections = []
        
        for section in sections:
            if section["estimated_tokens"] > self.max_tokens * 0.9:  # 接近或超过限制
                oversized_sections.append(section)
            else:
                normal_sections.append(section)
        
        # 处理正常大小的章节：尽量将相邻章节组合到接近上限
        for section in normal_sections:
            tokens = section["estimated_tokens"]
            
            # 如果添加当前章节会超过限制，先保存当前批次
            if current_tokens + tokens > self.max_tokens and current_batch:
                # 合并当前批次中的章节
                merged_section = self._merge_sections(current_batch)
                optimized_sections.append(merged_section)
                current_batch = []
                current_tokens = 0
            
            # 添加当前章节到批次
            current_batch.append(section)
            current_tokens += tokens
        
        # 处理最后一个批次
        if current_batch:
            merged_section = self._merge_sections(current_batch)
            optimized_sections.append(merged_section)
        
        # 处理超大章节，需要分割
        for section in oversized_sections:
            split_sections = self._split_large_section(section)
            optimized_sections.extend(split_sections)
        
        print(f"[INFO] 优化后的部分数量: {len(optimized_sections)}")
        return optimized_sections
    
    def _merge_sections(self, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """合并多个章节为一个部分"""
        if len(sections) == 1:
            return sections[0]
        
        # 合并标题和内容
        titles = [s["title"] for s in sections]
        merged_title = " + ".join(titles)
        
        content_parts = []
        for section in sections:
            content_parts.append(f"[{section['title']}]\n\n{section['content']}")
        
        merged_content = "\n\n---\n\n".join(content_parts)
        merged_tokens = sum(s["estimated_tokens"] for s in sections)
        
        return {
            "title": merged_title,
            "content": merged_content,
            "type": "merged_chapters",
            "estimated_tokens": merged_tokens,
            "merged_from": titles
        }
    
    def _split_large_section(self, section: Dict[str, Any]) -> List[Dict[str, Any]]:
        """将大型章节分割成多个部分，尽量按段落分割"""
        content = section["content"]
        estimated_tokens = section["estimated_tokens"]
        title = section["title"]
        
        # 计算需要分割的部分数
        parts_needed = math.ceil(estimated_tokens / (self.max_tokens * 0.9))
        
        # 按段落分割
        paragraphs = content.split("\n\n")
        
        # 如果段落太少，改为按句子分割
        if len(paragraphs) < parts_needed * 2:
            text_chunks = []
            current_chunk = ""
            current_tokens = 0
            target_tokens = estimated_tokens / parts_needed
            
            for para in paragraphs:
                para_tokens = self._estimate_tokens(para)
                
                # 如果当前段落非常大，需要进一步分割
                if para_tokens > target_tokens:
                    sentences = re.split(r'(?<=[.!?])\s+', para)
                    for sentence in sentences:
                        sent_tokens = self._estimate_tokens(sentence)
                        if current_tokens + sent_tokens > target_tokens and current_chunk:
                            text_chunks.append(current_chunk)
                            current_chunk = sentence
                            current_tokens = sent_tokens
                        else:
                            current_chunk += " " + sentence if current_chunk else sentence
                            current_tokens += sent_tokens
                else:
                    if current_tokens + para_tokens > target_tokens and current_chunk:
                        text_chunks.append(current_chunk)
                        current_chunk = para
                        current_tokens = para_tokens
                    else:
                        current_chunk += "\n\n" + para if current_chunk else para
                        current_tokens += para_tokens
            
            # 添加最后一个块
            if current_chunk:
                text_chunks.append(current_chunk)
        else:
            # 每部分的段落数
            paras_per_part = max(1, len(paragraphs) // parts_needed)
            
            # 分组
            text_chunks = []
            for i in range(0, len(paragraphs), paras_per_part):
                chunk = "\n\n".join(paragraphs[i:i+paras_per_part])
                text_chunks.append(chunk)
        
        # 创建分割后的部分
        split_sections = []
        for i, chunk in enumerate(text_chunks):
            split_sections.append({
                "title": f"{title} (Part {i+1}/{len(text_chunks)})",
                "content": chunk,
                "type": "split_chapter",
                "estimated_tokens": self._estimate_tokens(chunk),
                "original_chapter": title
            })
        
        return split_sections
    
    def _estimate_tokens(self, text: str) -> int:
        """
        估算文本中的token数量
        一般来说，英文中每个单词约为1.3个token，
        或者每4个字符约为1个token
        
        对于中文，通常每个字符约1个token
        """
        # 检测文本是否主要是中文
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        is_mainly_chinese = chinese_chars > len(text) * 0.5
        
        if is_mainly_chinese:
            # 中文文本：每个字符大约是1个token
            return int(len(text) * 0.75)  # 稍微保守一点
        else:
            # 英文文本
            words = len(re.findall(r'\b\w+\b', text))
            char_count = len(text)
            
            # 两种估算方法的平均值可能更准确
            word_based = int(words * 1.3)
            char_based = int(char_count * 0.25)
            
            return max(word_based, char_based)  # 取较大值以避免低估